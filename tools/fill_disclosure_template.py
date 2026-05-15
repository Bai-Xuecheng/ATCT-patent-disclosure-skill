#!/usr/bin/env python3
"""
Fill the user's patent disclosure DOCX template.

Usage:
  python3 tools/fill_disclosure_template.py \
    --template 专利交底书模板.docx \
    --payload payload.json \
    --output 案件名称-专利交底书.docx

Payload schema: see tools/example_disclosure_payload.json.
"""
from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
try:
    from PIL import Image
except Exception:  # Pillow is optional; python-docx can still insert supported image files.
    Image = None


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "、".join(str(x).strip() for x in value if str(x).strip())
    return str(value).replace("\r\n", "\n").replace("\r", "\n").strip()


def join_keywords(value: Any, sep: str = "、") -> str:
    if isinstance(value, list):
        return sep.join(str(x).strip() for x in value if str(x).strip())
    return clean_text(value)


def set_run_font(run, font_name: str = "宋体", size_pt: Optional[int] = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt:
        run.font.size = Pt(size_pt)


def clear_cell(cell) -> None:
    # Keep the first paragraph object where possible to preserve table structure.
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    # Ensure there is at least one paragraph after clearing.
    cell._tc.add_p()


def set_cell_text(cell, text: Any, font_name: str = "宋体", size_pt: int = 10, align: Optional[int] = None) -> None:
    text = clean_text(text)
    clear_cell(cell)
    paras = text.split("\n") if text else [""]
    first = True
    for para_text in paras:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(para_text)
        set_run_font(run, font_name, size_pt)
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass


def fitted_width_inches(image_path: str, max_width_inches: float, max_height_inches: Optional[float] = None) -> float:
    """Return a width that preserves aspect ratio and fits within optional max height."""
    if not max_height_inches or Image is None:
        return max_width_inches
    try:
        with Image.open(image_path) as im:
            w, h = im.size
        if not w or not h:
            return max_width_inches
        candidate_height = max_width_inches * (h / w)
        if candidate_height <= max_height_inches:
            return max_width_inches
        return max(1.0, max_height_inches / (h / w))
    except Exception:
        return max_width_inches


def add_fitted_picture(run, image_path: str, max_width_inches: float, max_height_inches: Optional[float] = None) -> None:
    width_inches = fitted_width_inches(image_path, max_width_inches, max_height_inches)
    run.add_picture(image_path, width=Inches(width_inches))


def add_picture_to_cell(cell, image_path: str, width_inches: float = 2.35, caption: str = "", max_height_inches: Optional[float] = None) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path and Path(image_path).exists():
        run = p.add_run()
        add_fitted_picture(run, image_path, width_inches, max_height_inches=max_height_inches)
        if caption:
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.add_run(caption)
            set_run_font(r, size_pt=9)
    else:
        r = p.add_run(caption or "")
        set_run_font(r, size_pt=10)


def row_texts(table, row_idx: int) -> List[str]:
    return [clean_text(c.text) for c in table.rows[row_idx].cells]


def find_last_cell_containing(table, row_idx: int, needle: str) -> Optional[int]:
    idx = None
    for i, txt in enumerate(row_texts(table, row_idx)):
        if needle in txt:
            idx = i
    return idx


def cell_after(table, row_idx: int, needle: str, fallback: int) -> Any:
    row = table.rows[row_idx]
    idx = find_last_cell_containing(table, row_idx, needle)
    if idx is None:
        target = fallback
    else:
        target = min(idx + 1, len(row.cells) - 1)
    return row.cells[target]


def safe_cell(table, row_idx: int, col_idx: int) -> Any:
    row = table.rows[row_idx]
    col_idx = max(0, min(col_idx, len(row.cells) - 1))
    return row.cells[col_idx]


def rel_path(path: str, base: Path) -> str:
    if not path:
        return ""
    p = Path(path)
    if p.is_absolute():
        return str(p)
    candidate = base / p
    return str(candidate if candidate.exists() else p)


def make_drawings_text(drawings: Iterable[Dict[str, Any]]) -> str:
    drawings = list(drawings or [])
    if not drawings:
        return ""
    nums = [clean_text(d.get("no")) or f"图{i+1}" for i, d in enumerate(drawings)]
    captions = []
    for i, d in enumerate(drawings):
        no = clean_text(d.get("no")) or f"图{i+1}"
        title = clean_text(d.get("title")) or "附图"
        captions.append(f"{no}为{title}")
    return "\n".join(nums + [""] + captions)


def fill_document(template_path: Path, payload: Dict[str, Any], output_path: Path) -> None:
    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("Template has no table; expected the provided patent disclosure template.")
    table = doc.tables[0]
    base_dir = Path(payload.get("_base_dir") or template_path.parent)

    meta = payload.get("meta", {}) or {}
    writer = meta.get("撰写人", {}) or {}
    contact = meta.get("技术问题联系人", {}) or {}

    # Header paragraphs: title and number.
    number = clean_text(meta.get("编号"))
    for p in doc.paragraphs:
        txt = clean_text(p.text)
        if txt.startswith("编号"):
            p.text = f"编号: {number}" if number else "编号: "
            for run in p.runs:
                set_run_font(run, size_pt=10)
            break

    # Rows 0-3: people/contact fields.
    set_cell_text(cell_after(table, 0, "撰写人", 4), writer.get("姓名", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 0, "手机", 7), writer.get("手机", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 0, "邮箱", 9), writer.get("邮箱", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 1, "座机", 7), writer.get("座机", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 1, "邮箱", 9), writer.get("邮箱", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    set_cell_text(cell_after(table, 2, "技术问题联系人", 4), contact.get("姓名", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 2, "手机", 7), contact.get("手机", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 2, "邮箱", 9), contact.get("邮箱", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 3, "座机", 7), contact.get("座机", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 3, "邮箱", 9), contact.get("邮箱", ""), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Core metadata.
    set_cell_text(cell_after(table, 4, "发明名称", 4), payload.get("title", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 5, "中文", 4), join_keywords(payload.get("keywords_cn")), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cell_after(table, 6, "英文", 4), join_keywords(payload.get("keywords_en"), sep="、"), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Search formula row. Preserve two cells: Chinese and English.
    cn_search = clean_text(payload.get("search_formula_cn"))
    en_search = clean_text(payload.get("search_formula_en"))
    if cn_search and not cn_search.startswith("中文"):
        cn_search = "中文：" + cn_search
    if en_search and not (en_search.startswith("英文") or en_search.startswith("English")):
        en_search = "英文：" + en_search
    search_cn_cell = cell_after(table, 7, "检索式", 3)
    set_cell_text(search_cn_cell, cn_search)
    # English cell: next physical cell after Chinese cell when possible.
    cn_idx = None
    for i, c in enumerate(table.rows[7].cells):
        if c is search_cn_cell:
            cn_idx = i
            break
    set_cell_text(safe_cell(table, 7, (cn_idx + 1) if cn_idx is not None else 4), en_search)

    # Prior art and relevance.
    set_cell_text(cell_after(table, 9, "期刊", 4), payload.get("prior_art_literature", ""))
    set_cell_text(cell_after(table, 10, "中国专利", 4), payload.get("prior_art_patents", ""))
    set_cell_text(cell_after(table, 11, "相关性分析", 2), payload.get("relevance_analysis", ""))

    # Abstract and representative figure.
    set_cell_text(safe_cell(table, 13, 1), payload.get("abstract", ""))
    rep = rel_path(clean_text(payload.get("representative_figure")), base_dir)
    if rep:
        add_picture_to_cell(safe_cell(table, 13, 6), rep, width_inches=2.35, caption="代表性附图", max_height_inches=4.8)

    # Main content.
    set_cell_text(safe_cell(table, 15, 0), payload.get("technical_field", ""))
    set_cell_text(safe_cell(table, 17, 0), payload.get("background", ""))
    set_cell_text(cell_after(table, 19, "有益的技术效果", 3), payload.get("purpose_and_effect", ""))
    set_cell_text(cell_after(table, 20, "所有的实施方式", 3), payload.get("embodiments", ""))
    set_cell_text(cell_after(table, 21, "可替代的技术方案", 3), payload.get("alternatives", "无") or "无")

    drawings = payload.get("drawings", []) or []
    drawing_cell = safe_cell(table, 23, 0)
    clear_cell(drawing_cell)
    # If images exist, rely on image captions to avoid duplicate or awkward split captions.
    has_rendered_images = False
    for d in drawings:
        path = rel_path(clean_text(d.get("path")), base_dir)
        if path and Path(path).exists():
            has_rendered_images = True
            break
    p = drawing_cell.paragraphs[0]
    r = p.add_run("" if has_rendered_images else make_drawings_text(drawings))
    set_run_font(r, size_pt=10)
    # Then optionally add pictures with captions.
    for i, d in enumerate(drawings):
        path = rel_path(clean_text(d.get("path")), base_dir)
        no = clean_text(d.get("no")) or f"图{i+1}"
        title = clean_text(d.get("title")) or "附图"
        if path and Path(path).exists():
            pp = drawing_cell.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pp.add_run()
            add_fitted_picture(run, path, max_width_inches=5.2, max_height_inches=5.7)
            cp = drawing_cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = cp.add_run(f"{no}为{title}")
            set_run_font(rr, size_pt=9)

    set_cell_text(safe_cell(table, 25, 0), payload.get("protection_points", ""))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True, help="Path to 专利交底书模板.docx")
    parser.add_argument("--payload", required=True, help="Path to payload JSON")
    parser.add_argument("--output", required=True, help="Output DOCX path")
    args = parser.parse_args()

    template_path = Path(args.template).expanduser().resolve()
    payload_path = Path(args.payload).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    with payload_path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    payload.setdefault("_base_dir", str(payload_path.parent))

    fill_document(template_path, payload, output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
